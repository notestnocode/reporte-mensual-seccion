import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

# --- 1. CONFIGURACIÓN DEL SISTEMA (PROMPT MAESTRO) ---
SYSTEM_PROMPT = """Actúa como el Asistente Digital del Grupo Scout 19 Paxtu. Tu objetivo es generar el "Reporte Mensual de Sección" mediante una entrevista con el Scouter.

1. DINÁMICA DE TRABAJO:
- Entrevista al Scouter de la sección (el reporte es para una sola sección).
- Si el Scouter te da datos narrativos, extráelos y clasifícalos en la tabla correspondiente.
- Al finalizar la recolección, presenta un resumen de la información recibida, espera confirmación y posteriormente genera el reporte completo.

2. ESTRUCTURA DEL REPORTE (FORMATO PARA COPIAR A WORD/GOOGLE DOCS):

2.1 TITULO
# GRUPO 19 PAXTU - REPORTE DE SECCIÓN [Sección]
## Mes: [Mes/Año] Elabora: [Nombre]

2.2 ACTIVIDADES REALIZADAS
### ACTIVIDADES
| Fecha | Tipo de Actividad | Asistencia (Jovenes/Adultos) | Descripción | Evaluación |
| :--- | :--- | :--- | :--- | :--- |

2.3. MEMBRESÍA
### MEMBRESÍA
| Total Miembros | Registrados | Prospectos | Altas | Bajas |
| :--- | :--- | :--- | :--- | :--- |
- **Lista de Altas:** [Nombres]
- **Lista de Bajas:** [Nombres]

IV. FINANZAS (CAJA CHICA)
### FINANZAS
| Saldo Inicial | Total Ingresos | Total Egresos | Saldo Final |
| :--- | :--- | :--- | :--- |
### Detalle de movimientos:
[Lista]

V. RESUMEN DE PROGRESIÓN (CONTEO)
### PROGRESIONES ENTREGADAS
| Nombre de la Insignia | Cantidad Total |
| :--- | :--- |

VI. DETALLE DE PROGRESIÓN (INDIVIDUAL)
### PROGRESIONES PERSONALES
| Tipo de Insignia | Nombre de la Insignia | Fecha de Entrega | Nombre del Joven |
| :--- | :--- | :--- | :--- |

VII. ASUNTOS PARA LLEVAR A CONSEJO
### PUNTOS Y ACUERDOS PARA PRESENTAR EN CONSEJO
| Prioridad | Observación / Solicitud | Estatus |
| :--- | :--- | :--- |

3. REGLAS CRÍTICAS:
- NO uses bloques de código (fondo gris/backticks). Entrega el texto y las tablas directamente en el chat.
- No inventes datos. Si tienes duda pregunta discretamente. Si una tabla no tiene información, llénala con "Sin movimientos este mes".
- Si se menciona una entrega de insignia en la narrativa, regístrala automáticamente en las dos tablas de Progresión.
- Usa fuentes en negrita y títulos claros."""

# --- FUNCIONES PARA WORD (TABLAS MORADAS) ---
def set_cell_background(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generar_docx(texto_reporte):
    doc = Document()
    MORADO_SCOUT = "4A267A" 
    lineas = texto_reporte.split('\n')
    i = 0
    while i < len(lineas):
        linea = lineas[i].strip()
        if not linea:
            i += 1
            continue
        if linea.startswith('# '):
            doc.add_heading(linea.replace('# ', ''), level=1)
        elif linea.startswith('|') and i + 1 < len(lineas) and '| :---' in lineas[i+1]:
            columnas = [c.strip() for c in linea.split('|') if c.strip()]
            tabla = doc.add_table(rows=1, cols=len(columnas))
            tabla.style = 'Table Grid'
            hdr_cells = tabla.rows[0].cells
            for idx, col_name in enumerate(columnas):
                hdr_cells[idx].text = col_name
                set_cell_background(hdr_cells[idx], MORADO_SCOUT)
                run = hdr_cells[idx].paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
            i += 2
            while i < len(lineas) and lineas[i].strip().startswith('|'):
                datos = [d.strip() for d in lineas[i].split('|') if d.strip()]
                if datos:
                    row_cells = tabla.add_row().cells
                    for idx, valor in enumerate(datos):
                        if idx < len(row_cells):
                            row_cells[idx].text = valor
                i += 1
            continue
        else:
            doc.add_paragraph(linea)
        i += 1
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 2. CONFIGURACIÓN DE PÁGINA E INICIALIZACIÓN ---
st.set_page_config(page_title="Reporte Paxtu 19", page_icon="⚜️")
st.title("🤖 Asistente de Reportes - Grupo 19 Paxtu")

if "messages" not in st.session_state:
    st.session_state.messages = []
if "reporte_listo" not in st.session_state:
    st.session_state.reporte_listo = False
if "ultimo_reporte" not in st.session_state:
    st.session_state.ultimo_reporte = ""

# --- 3. BARRA LATERAL ---
with st.sidebar:
    st.header("📋 Guía para el Scouter")
    st.markdown("""
    **💡 Tip de Dictado:**
    Usa el **micrófono de tu teclado** en el celular. ¡Es mucho más rápido!
    ---
    **🗣️ Ejemplo de cómo hablar:**
    > *"Soy Baloo, reporte de Manada de octubre. El día 12 fuimos al parque con 10 lobatos. Le entregamos la insignia de 'Rastreador' a Juan Pérez. Gastamos $150 en dulces."*
    
    ---
    **Secciones que incluye tu reporte:**
    1. **Encabezado:** Título y responsable.
    2. **Actividades:** Fechas y evaluación.
    3. **Membresía:** Altas, bajas y registros.
    4. **Finanzas:** Movimientos de caja chica.
    5. **Resumen Progresión:** Conteo total.
    6. **Detalle Progresión:** Quién recibió qué.
    7. **Asuntos de Consejo:** Avisos para el Grupo.
    
    ---
    **Proceso:**
    Confirma los datos cuando el bot te dé el resumen y descarga tu Word con tablas en morado.
    """)
    
    st.divider()
    if st.button("🗑️ Nuevo Reporte"):
        st.session_state.messages = []
        st.session_state.reporte_listo = False
        st.session_state.ultimo_reporte = ""
        st.rerun()

# --- 4. CONEXIÓN API ---
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
model = genai.GenerativeModel(model_name='gemini-2.5-flash', system_instruction=SYSTEM_PROMPT)

# Mostrar historial
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# --- 5. LÓGICA DE INTERACCIÓN ---
if prompt := st.chat_input("Cuéntame sobre el mes de la sección..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    try:
        history_google = [{"role": "user" if m["role"] == "user" else "model", "parts": [m["content"]]} for m in st.session_state.messages[:-1]]
        chat = model.start_chat(history=history_google)
        response = chat.send_message(prompt)
        
        with st.chat_message("assistant"):
            st.markdown(response.text)
            # Detección flexible del reporte
            if "# GRUPO 19 PAXTU" in response.text or "| Fecha |" in response.text:
                st.session_state.reporte_listo = True
                st.session_state.ultimo_reporte = response.text
        
        st.session_state.messages.append({"role": "assistant", "content": response.text})
        st.rerun() # Forzamos recarga para que el botón aparezca inmediatamente

    except Exception as e:
        st.error(f"Error: {str(e)}")

# --- 6. MOSTRAR BOTÓN DE DESCARGA (PERSISTENTE) ---
if st.session_state.reporte_listo:
    st.divider()
    st.success("✅ ¡Reporte detectado! Ya puedes descargarlo.")
    archivo_word = generar_docx(st.session_state.ultimo_reporte)
    st.download_button(
        label="📥 Descargar Reporte para Word (.docx)",
        data=archivo_word,
        file_name="Reporte_Paxtu19.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="btn_descarga"
    )
